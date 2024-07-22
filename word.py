from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Document
doc = Document()

# Cover Page
cover_page = doc.add_section(0)
cover_page_start_type = cover_page.start_type
cover_page.page_height
cover_page.page_width

cover_title = doc.add_heading('รายงานวิจัยตลาด: ถุงขยะบน Lazada', level=1)
cover_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph()

subtitle = doc.add_heading('โดย [ชื่อคุณ]', level=2)
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph()

date = doc.add_heading('วันที่ [วันที่]', level=3)
date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_page_break()

# Table of Contents
toc_heading = doc.add_heading('สารบัญ', level=1)

paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar3)

doc.add_page_break()

# Market Research Report Content
doc.add_heading('เกณฑ์สำหรับการวิเคราะห์คู่แข่ง', level=1)

doc.add_heading('ข้อมูลบริษัท', level=2)
doc.add_heading('ชื่อบริษัท', level=3)
doc.add_heading('ชื่อแบรนด์', level=3)

doc.add_heading('ค่าใช้จ่าย', level=2)
doc.add_heading('ราคา', level=3)
doc.add_heading('ค่าจัดส่ง', level=3)

doc.add_heading('การจัดส่ง', level=2)
doc.add_heading('เวลาจัดส่ง', level=3)

doc.add_heading('บรรจุภัณฑ์และการติดฉลาก', level=2)
doc.add_heading('บรรจุภัณฑ์', level=3)
doc.add_heading('ฉลาก', level=3)
doc.add_heading('รูปภาพ', level=3)
doc.add_heading('วิธีการพับถุง (การพับถุง)', level=3)

doc.add_heading('คุณภาพของสินค้า', level=2)
doc.add_heading('คุณภาพของวัสดุถุงขยะ', level=3)
doc.add_heading('ขนาดและความจุ', level=3)
doc.add_heading('คุณสมบัติเพิ่มเติม', level=3)

doc.add_heading('การประเมินมูลค่า', level=2)
doc.add_heading('ความคุ้มค่าของเงิน', level=3)

doc.add_heading('ความคิดเห็นเพิ่มเติม', level=2)
doc.add_paragraph('[ข้อมูลที่เกี่ยวข้องอื่นๆ]')

doc.add_heading('การวิเคราะห์คู่แข่ง', level=1)

competitors = [
    "แพนด้า (ที่เกี่ยวข้อง, Best seller)",
    "ม.บก (ที่เกี่ยวข้อง)",
    "ถุงขยะราคาถูกจากโรงงาน MUTU shop (ที่เกี่ยวข้อง, Best seller)",
    "ดวงเฮง (ที่เกี่ยวข้อง)",
    "พญาขิง (ที่เกี่ยวข้อง)",
    "ถุงขยะต้านหนา ราคาส่ง (Here packing) (ที่เกี่ยวข้อง)",
    "ถุงขยะอันดับ 1 (Best seller)",
    "จัมโบ้ (Best seller)",
    "ถุงขยะยี่ห้อ Easy life (Best seller)",
    "ถุงขยะต้า หนาพิเศษ (Best Q) (การจัดอันดับ)",
    "ถุงขยะต้า (Prompt999) (เสร็จรับเจลล้าง)",
    "ถุงขยะถุง (ชื่อจะถูกกำหนด)"
]

for competitor in competitors:
    doc.add_heading(f'คู่แข่ง: {competitor}', level=2)
    doc.add_heading('ข้อมูลบริษัท', level=3)
    doc.add_paragraph(f'ชื่อบริษัท: {competitor.split()[0]}')
    doc.add_paragraph('ชื่อแบรนด์: [ชื่อแบรนด์]')
    doc.add_heading('ค่าใช้จ่าย', level=3)
    doc.add_paragraph('ราคา: [รายละเอียดราคา]')
    doc.add_paragraph('ค่าจัดส่ง: [รายละเอียด]')
    doc.add_heading('การจัดส่ง', level=3)
    doc.add_paragraph('เวลาจัดส่ง: [รายละเอียด]')
    doc.add_heading('บรรจุภัณฑ์และการติดฉลาก', level=3)
    doc.add_paragraph('บรรจุภัณฑ์: [รายละเอียด]')
    doc.add_paragraph('ฉลาก: [รายละเอียด]')
    doc.add_paragraph('รูปภาพ: [ลิงก์รูปภาพหรือคำอธิบาย]')
    doc.add_paragraph('วิธีการพับถุง (การพับถุง): [รายละเอียด]')
    doc.add_heading('คุณภาพของสินค้า', level=3)
    doc.add_paragraph('คุณภาพของวัสดุถุงขยะ: [รายละเอียด]')
    doc.add_paragraph('ขนาดและความจุ: [รายละเอียด]')
    doc.add_paragraph('คุณสมบัติเพิ่มเติม: [รายละเอียด]')
    doc.add_heading('การประเมินมูลค่า', level=3)
    doc.add_paragraph('ความคุ้มค่าของเงิน: [การประเมิน]')
    doc.add_heading('ความคิดเห็นเพิ่มเติม', level=3)
    doc.add_paragraph('[ข้อมูลที่เกี่ยวข้องอื่นๆ]')

doc.add_heading('สรุปและข้อเสนอแนะ', level=1)
doc.add_paragraph('สรุปผลการวิจัย')
doc.add_paragraph('ข้อเสนอแนะตามการวิเคราะห์')

# Save the document
doc.save('Market_Research_Report.docx')
